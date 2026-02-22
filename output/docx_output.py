"""
Word Document Output Generator
-------------------------------
Produces a .docx file containing all citations with their resolved URLs,
confidence levels, and claims. Formatted with clear headings and color-coded
confidence indicators for easy review.
"""

from pathlib import Path
from datetime import datetime
from typing import List, Dict


CONFIDENCE_COLORS = {
    "HIGH":   "375623",   # dark green
    "MEDIUM": "7B5E00",   # dark amber
    "LOW":    "9C1A1A",   # dark red
}

CONFIDENCE_BG = {
    "HIGH":   "D4EDDA",
    "MEDIUM": "FFF3CD",
    "LOW":    "F8D7DA",
}

TYPE_LABELS = {
    "article":    "Journal Article",
    "book":       "Book",
    "report":     "Corporate Report",
    "government": "Government Document",
    "web":        "Web Source",
    "unknown":    "Unknown",
}


def generate_docx(resolved: List[Dict], output_path: Path, source_doc: str = "") -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    doc = Document()

    # ---- Document title ----
    title = doc.add_heading("Citation Checker Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Source: {source_doc}  |  Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ---- Summary table ----
    high   = sum(1 for r in resolved if r["confidence"] == "HIGH")
    medium = sum(1 for r in resolved if r["confidence"] == "MEDIUM")
    low    = sum(1 for r in resolved if r["confidence"] == "LOW")

    summary_heading = doc.add_paragraph()
    summary_heading.add_run("Summary").bold = True

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["Total Citations", "High Confidence", "Medium Confidence", "Low Confidence"]
    values  = [str(len(resolved)), str(high), str(medium), str(low)]
    for i, (h, v) in enumerate(zip(headers, values)):
        table.cell(0, i).text = h
        table.cell(1, i).text = v
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # ---- One section per citation ----
    for idx, citation in enumerate(resolved, 1):
        confidence = citation.get("confidence", "LOW")
        key        = citation.get("key", "Unknown")
        title_str  = citation.get("title", "")
        raw        = citation.get("raw", "")
        claims     = citation.get("claims", [])
        urls       = citation.get("urls", [])
        notes      = citation.get("notes", "")
        ctype      = citation.get("type", "unknown")

        # Citation heading with confidence indicator
        heading_text = f"{idx}. {key}  [{confidence}]"
        h = doc.add_heading(heading_text, level=2)
        for run in h.runs:
            color_hex = CONFIDENCE_COLORS.get(confidence, "000000")
            r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)

        # Apply background shading to the heading paragraph
        _set_paragraph_bg(h, CONFIDENCE_BG.get(confidence, "FFFFFF"))

        # Type pill
        type_p = doc.add_paragraph()
        type_run = type_p.add_run(f"Type: {TYPE_LABELS.get(ctype, ctype)}")
        type_run.font.size = Pt(9)
        type_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        type_run.italic = True

        # Full citation
        doc.add_paragraph("Full Citation:", style="Intense Quote") \
           if "Intense Quote" in [s.name for s in doc.styles] \
           else doc.add_paragraph("Full Citation:")
        raw_p = doc.add_paragraph(raw)
        raw_p.paragraph_format.left_indent = Inches(0.3)
        for run in raw_p.runs:
            run.font.size = Pt(10)

        # Claims
        if claims:
            label = doc.add_paragraph()
            label.add_run("Claims Found in Document:").bold = True
            for claim in claims[:3]:
                cp = doc.add_paragraph(f'"{claim}"')
                cp.paragraph_format.left_indent = Inches(0.3)
                for run in cp.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True

        # URLs
        if urls:
            label = doc.add_paragraph()
            label.add_run("Sources Found:").bold = True
            for u in urls:
                url_p = doc.add_paragraph()
                url_p.paragraph_format.left_indent = Inches(0.3)
                # Add hyperlink
                _add_hyperlink(url_p, u["url"], u.get("title") or u["url"])
                snippet = u.get("snippet", "")
                if snippet:
                    snip_run = url_p.add_run(f"\n  {snippet[:160]}")
                    snip_run.font.size = Pt(8)
                    snip_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Notes / confidence explanation
        if notes:
            note_p = doc.add_paragraph(f"Note: {notes}")
            note_p.paragraph_format.left_indent = Inches(0.3)
            for run in note_p.runs:
                run.font.size = Pt(9)
                color_hex = CONFIDENCE_COLORS.get(confidence, "555555")
                r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

        doc.add_paragraph()  # spacer between citations

    doc.save(output_path)


def _set_paragraph_bg(paragraph, hex_color: str):
    """Apply a background shading color to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink into a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style as blue underlined link
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    w_t = OxmlElement("w:t")
    w_t.text = text
    new_run.append(w_t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
